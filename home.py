import os
import json
import streamlit as st
import numpy as np
import tensorflow as tf
from groq import Groq
from docx import Document
from docx.shared import Inches
import tempfile
from PIL import Image


def Home(lang_code):
    # Translations Dictionary
    translations = {
        "en": {
            "title": "🌱 Welcome to Tomato Detection System!",
            "subheader": "Your ultimate companion for tomato planting, care, and disease management.",
            "features": """  
                Use the Tomato System to:  
                - Detect and diagnose diseases in tomato plants.  
                - Get expert advice and treatment options with Planty AI.  
                - Shop for high-quality fertilizers, pesticides, and seeds.  
            """,
            "offer": "🌟 Special Offer for Tomato Lovers!",
            "buy_seeds": "Buy Premium Tomato Seeds Now!",
            "seed_benefits": """  
                - High yield and disease resistance.  
                - Suitable for all climates.  
                - Special 20% discount for a limited time!  
            """,
            "why_choose": """  
                🌟 Why choose our seeds?  
                - Tested and trusted by farmers worldwide.  
                - Supports sustainable and organic farming practices.  
                - Guaranteed freshness and germination rates.  
            """,
            # "footer_title" : "## 🍃 Thank You for Using Our System!",
            "footer_text" : "We are committed to helping farmers and gardeners maintain healthy tomato crops.  \nIf you have any questions or need further assistance, feel free to reach out.  \n**Happy farming! 🌱**",
            "shop": "👉 Visit Our Shop",
            "disease_diagnosis": "Tomato Disease Diagnosis",
            "upload_prompt": "📤 Upload an image of the tomato leaf",
            "invalid_image": "❌ Invalid Image",
            "low_quality" : "### ⚠️ Low Image Quality",
            "upload_clear" : "Please upload a clearer image for better diagnosis.",
            "upload_valid_image": "Please upload a valid image of a tomato leaf for diagnosis.",
            "disease_detected": "✅ Disease Detected:",
            "confidence": "Confidence:",
            "solution_info": "### Information and Solution:",
            "generate_pdf": "📄 Generate PDF Report",
            "download_report": "📥 Download Diagnosis Report",
            "upload_image": "Uploaded Image",
            "Bacterial_spot" : "Bacterial Spot",
            "Early_blight" : "Early Blight",
            "Late_blight" : "Late Blight",
            "Leaf_Mold" : "Leaf Mold",
            "No_tomato_leaf" : "No Tomato Leaf",
            "Septoria_leaf_spot" : "Septoria Leaf Spot",
            "Spider_mites_Two-spotted_spider_mite" : "Spider Mites (Two-Spotted Spider Mite)",
            "Target_Spot" : "Target Spot",
            "Tomato_Yellow_Leaf_Curl_Virus" : "Tomato Yellow Leaf Curl Virus",
            "Tomato_mosaic_virus" : "Tomato Mosaic Virus",
            "Healthy" : "Healthy",
            "powdery_mildew" : "Powdery Mildew",
            "disease_info" : "**Information:**",
            "treatment_solutions" : "**Treatment Solutions:**",
            "pesticide_recommendations" : "**Pesticide Recommendations:**",
            "provide_info" : "Provide information, treatment solutions, and pesticide recommendations for {predicted_class} in tomatoes.",
            "footer_title" : "🍃 Thank You for Using Our System!",
            "footer_description" : """
            We are dedicated to helping farmers and gardeners maintain **healthy tomato crops**.  
            If you have any questions or need assistance, feel free to reach out.  
            """,
            "footer_closing" : "🌱 **Happy Farming!**",
            "generate_doc" : "📄 Generate Word Report",
            "download_report1" : "📥 Download Diagnosis Report (Word)",
            "title_doc" : "🍅 Tomato Disease Diagnosis Report",

        },
        "hi": {
            "title": "🌱 टमाटर पहचान प्रणाली में आपका स्वागत है!",
            "subheader": "टमाटर की खेती, देखभाल और रोग प्रबंधन के लिए आपका सबसे अच्छा साथी।",
            "features": """  
                टमाटर सिस्टम का उपयोग करें:  
                - टमाटर के पौधों में बीमारियों का पता लगाएं और उनका निदान करें।  
                - Planty AI के साथ विशेषज्ञ सलाह और उपचार विकल्प प्राप्त करें।  
                - उच्च गुणवत्ता वाले उर्वरकों, कीटनाशकों और बीजों की खरीदारी करें।  
            """,
            "offer": "🌟 टमाटर प्रेमियों के लिए विशेष ऑफर!",
            "buy_seeds": "प्रीमियम टमाटर के बीज अभी खरीदें!",
            "seed_benefits": """  
                - उच्च उपज और रोग प्रतिरोधक क्षमता।  
                - सभी जलवायु के लिए उपयुक्त।  
                - सीमित समय के लिए विशेष 20% छूट!  
            """,
            "why_choose": """  
                🌟 हमारे बीज क्यों चुनें?  
                - दुनिया भर के किसानों द्वारा परीक्षण और भरोसेमंद।  
                - टिकाऊ और जैविक खेती का समर्थन करता है।  
                - ताजगी और अंकुरण दर की गारंटी।  
            """,
            # "footer_title" : "## 🍃 हमारे सिस्टम का उपयोग करने के लिए धन्यवाद!",
            "footer_text" : "हम किसानों और बागवानों को स्वस्थ टमाटर की फसल बनाए रखने में मदद करने के लिए प्रतिबद्ध हैं।  \nयदि आपके कोई प्रश्न हैं या अधिक सहायता की आवश्यकता है, तो बेझिझक हमसे संपर्क करें।  \n**खुशहाल खेती करें! 🌱**",
            "shop": "👉 हमारी दुकान पर जाएं",
            "disease_diagnosis": "टमाटर रोग निदान",
            "upload_prompt": "📤 टमाटर की पत्ती की छवि अपलोड करें",
            "invalid_image": "❌ अमान्य छवि",
            "low_quality" : "### ⚠️ कम छवि गुणवत्ता",
            "upload_clear" : "बेहतर निदान के लिए कृपया एक स्पष्ट छवि अपलोड करें।",
            "upload_valid_image": "कृपया निदान के लिए एक मान्य टमाटर पत्ती की छवि अपलोड करें।",
            "disease_detected": "✅ पता चला रोग:",
            "confidence": "विश्वास स्तर:",
            "solution_info": "### जानकारी और समाधान:",
            "generate_pdf": "📄 पीडीएफ रिपोर्ट बनाएं",
            "download_report": "📥 निदान रिपोर्ट डाउनलोड करें",
            "upload_image": "अपलोड की गई छवि",
            "Bacterial_spot" : "बैक्टीरियल स्पॉट",
            "Early_blight" : "आरंभिक झुलसा",
            "Late_blight" : "देर से झुलसा",
            "Leaf_Mold" : "पत्तों का फफूंद",
            "No_tomato_leaf" : "कोई टमाटर पत्ता नहीं",
            "Septoria_leaf_spot" : "सेप्टोरिया पत्ती धब्बा",
            "Spider_mites_Two-spotted_spider_mite" : "स्पाइडर माइट्स (टू-स्पॉटेड स्पाइडर माइट)",
            "Target_Spot" : "टार्गेट स्पॉट",
            "Tomato_Yellow_Leaf_Curl_Virus" : "टमाटर पीला पत्ता कर्ल वायरस",
            "Tomato_mosaic_virus" : "टमाटर मोज़ेक वायरस",
            "Healthy" : "स्वस्थ",
            "powdery_mildew" : "पाउडरी मिल्ड्यू",
            "disease_info" : "**जानकारी:**",
            "treatment_solutions" : "**उपचार समाधान:**",
            "pesticide_recommendations" : "**कीटनाशक सिफारिशें:**",
            "provide_info" : "{predicted_class} के लिए टमाटरों में विस्तृत जानकारी, उपचार समाधान, और अनुशंसित कीटनाशकों की जानकारी प्रदान करें। इसमें रोग के कारण, लक्षण, रोकथाम के उपाय, उपचार विधियां (जैविक और रासायनिक दोनों), और उपयोगी कृषि तकनीकों का उल्लेख करें।",
            "description" : """
            **टमाटर पत्ती रोग पहचान प्रणाली** में आपका स्वागत है!  
            यह टूल किसानों और बागवानों को **टमाटर पत्तियों के रोगों की पहचान और निदान** करने में मदद करता है।  
            बस **अपने टमाटर के पत्ते की एक तस्वीर अपलोड करें** और तुरंत निदान और उपचार प्राप्त करें।
            """,
            "footer_title" : "🍃 हमारी प्रणाली का उपयोग करने के लिए धन्यवाद!",
            "footer_description" : """
            हम किसानों और बागवानों को **स्वस्थ टमाटर की फसल बनाए रखने में** मदद करने के लिए प्रतिबद्ध हैं।  
            यदि आपके कोई प्रश्न हैं या सहायता की आवश्यकता है, तो कृपया हमसे संपर्क करें।  
            """,
            "footer_closing" : "🌱 **खुशहाल खेती करें!**",
            "generate_doc" : "📄 वर्ड रिपोर्ट बनाएं",
            "download_report1" : "📥 निदान रिपोर्ट डाउनलोड करें (Word)",
            "title_doc" : "टमाटर रोग निदान रिपोर्ट",
        },
        "mr": {
            "title": "🌱 टोमॅटो शोध प्रणालीमध्ये आपले स्वागत आहे!",
            "subheader": "टोमॅटो लागवड, देखभाल आणि रोग व्यवस्थापनासाठी आपला सर्वोत्तम साथीदार.",
            "features": """  
                टोमॅटो सिस्टम वापरा:  
                - टोमॅटोच्या झाडांमधील रोग शोधा आणि त्यांचे निदान करा.  
                - Planty AI सोबत तज्ज्ञ सल्ला आणि उपचार पर्याय मिळवा.  
                - उच्च-गुणवत्तेची खतं, कीटकनाशके आणि बियाणे खरेदी करा.  
            """,
            "offer": "🌟 टोमॅटो प्रेमींसाठी विशेष ऑफर!",
            "buy_seeds": "प्रीमियम टोमॅटो बियाणे आत्ताच खरेदी करा!",
            "seed_benefits": """  
                - उच्च उत्पादन आणि रोग प्रतिकारशक्ती.  
                - सर्व हवामानासाठी उपयुक्त.  
                - मर्यादित वेळेसाठी 20% विशेष सूट!  
            """,
            "why_choose": """  
                🌟 आमची बियाणे का निवडायची?  
                - जगभरातील शेतकऱ्यांनी तपासलेली आणि विश्वासार्ह.  
                - टिकाऊ आणि सेंद्रिय शेतीला समर्थन देते.  
                - ताजेपणा आणि उगवणक्षमता हमी.  
            """,
            # "footer_title" : "## 🍃 आमच्या प्रणालीचा वापर केल्याबद्दल धन्यवाद!",
            "footer_text" : "शेतकरी आणि बागायतदार यांना निरोगी टोमॅटो पिके टिकवण्यासाठी मदत करण्यासाठी आम्ही वचनबद्ध आहोत.  \nतुम्हाला काही प्रश्न असल्यास किंवा अधिक मदतीची आवश्यकता असल्यास, कृपया आमच्याशी संपर्क साधा.  \n**सुखद शेती करा! 🌱**",
            "shop": "👉 आमच्या दुकानाला भेट द्या",
            "disease_diagnosis": "🍅 टोमॅटो रोग निदान",
            "upload_prompt": "📤 टोमॅटो पानाची प्रतिमा अपलोड करा",
            "invalid_image": "❌ अवैध प्रतिमा",
            "low_quality" : "### ⚠️ कमी प्रतिमा गुणवत्ता",
            "upload_clear" : "उत्तम निदानासाठी कृपया स्पष्ट प्रतिमा अपलोड करा.",
            "upload_valid_image": "कृपया निदानासाठी योग्य टोमॅटो पानाची प्रतिमा अपलोड करा.",
            "disease_detected": "✅ शोधलेला रोग:",
            "confidence": "विश्वास पातळी:",
            "solution_info": "### माहिती आणि उपाय:",
            "generate_pdf": "📄 पीडीएफ अहवाल तयार करा",
            "download_report": "📥 निदान अहवाल डाउनलोड करा",
            "upload_image": "अपलोड केलेली प्रतिमा",
            "Bacterial_spot" : "बॅक्टेरियल स्पॉट",
            "Early_blight" : "लवकर करपण",
            "Late_blight" : "उशीरा करपण",
            "Leaf_Mold" : "पाने बुरशी",
            "No_tomato_leaf" : "टोमॅटो पान नाही",
            "Septoria_leaf_spot" : "सेप्टोरिया पान डाग",
            "Spider_mites_Two-spotted_spider_mite" : "स्पायडर माइट्स (टू-स्पॉटेड स्पायडर माइट)",
            "Target_Spot" : "टार्गेट स्पॉट",
            "Tomato_Yellow_Leaf_Curl_Virus" : "टोमॅटो यलो लीफ कर्ल व्हायरस",
            "Tomato_mosaic_virus" : "टोमॅटो मोज़ेक व्हायरस",
            "Healthy" : "आरोग्यदायी",
            "powdery_mildew" : "पावडरी मिल्ड्यू",
            "disease_info" : "**माहिती:**",
            "treatment_solutions" : "**उपचार उपाय:**",
            "pesticide_recommendations" : "**कीटकनाशक शिफारसी:**",
            "provide_info" : "{predicted_class} साठी टोमॅटोवरील सविस्तर माहिती, उपचार उपाय, आणि शिफारस केलेले कीटकनाशक यांची माहिती प्रदान करा. यामध्ये रोगाची कारणे, लक्षणे, प्रतिबंधात्मक उपाय, उपचार पद्धती (सेंद्रिय आणि रासायनिक दोन्ही), तसेच उपयुक्त शेती तंत्रज्ञानाचा समावेश करा.",
            "description" : """
            **टोमॅटो पानांच्या रोगांचे निदान प्रणाली** मध्ये आपले स्वागत आहे!  
            हे साधन शेतकरी आणि माळी यांना **टोमॅटो पानांवरील रोग ओळखण्यास आणि निदान करण्यास** मदत करते.  
            फक्त **आपल्या टोमॅटोच्या पानाचा फोटो अपलोड करा** आणि त्वरित निदान व उपचार मिळवा.
            """,
            "footer_title" : "🍃 आमची प्रणाली वापरल्याबद्दल धन्यवाद!",
            "footer_description" : """
            आम्ही शेतकरी आणि माळी यांना **आरोग्यदायी टोमॅटो पीक राखण्यास** मदत करण्यासाठी वचनबद्ध आहोत.  
            तुम्हाला कोणतेही प्रश्न असल्यास किंवा अधिक मदतीची आवश्यकता असल्यास, कृपया आमच्याशी संपर्क साधा.  
            """,
            "footer_closing" : "🌱 **आनंदी शेती करा!**",
            "generate_doc" : "📄 वर्ड अहवाल तयार करा",
            "download_report1" : "📥 निदान अहवाल डाउनलोड करा (Word)",
            "title_doc" : "🍅 टोमॅटो रोग निदान अहवाल"
        }
    }

    t = translations[lang_code]

    # Header
    st.title ( t["title"] )
    st.subheader ( t["subheader"] )
    st.write ( t["features"] )

    st.markdown("---")
    
    model = tf.keras.models.load_model('tomato_disease_model.h5')
    
    # Classes for diseases
    classes = ['Bacterial_spot', 'Early_blight', 'Late_blight', 'Leaf_Mold', 'No_tomato_leaf', 'Septoria_leaf_spot',
               'Spider_mites Two-spotted_spider_mite', 'Target_Spot', 'Tomato_Yellow_Leaf_Curl_Virus',
               'Tomato_mosaic_virus', 'Healthy', 'powdery_mildew']

    # Function to preprocess image
    def preprocess_image(img) :
        img = img.resize ( (128, 128) )
        img_array = np.array ( img )
        img_array = np.expand_dims ( img_array, axis=0 )
        img_array = img_array / 255.0
        return img_array

    # Tomato Disease Diagnosis Section
    st.title(t["disease_diagnosis"])
    uploaded_file = st.file_uploader(t["upload_prompt"], type=["jpg", "jpeg", "png"])

    if uploaded_file:
        st.image(uploaded_file, caption=t['upload_image'], width=400)

        image = Image.open ( uploaded_file )
        processed_image = preprocess_image ( image )
        prediction = model.predict ( processed_image )

        predicted_class_index = np.argmax ( prediction, axis=1 )[0]
        predicted_class = classes[predicted_class_index]
        confidence = np.max ( prediction ) * 100

        try:
            # Configure Groq API key
            working_dir = os.path.dirname(os.path.abspath(__file__))
            config_data = json.load(open(f"{working_dir}/config.json"))
            GROQ_API_KEY = config_data["GROQ_API_KEY"]
            os.environ["GROQ_API_KEY"] = GROQ_API_KEY

            client = Groq()

            # Query Groq
            messages = [
                {"role": "system", "content": "You are an expert in tomato diseases and treatment solutions."},
                {"role": "user", "content": f"{t['provide_info'].format(predicted_class=predicted_class)}"},

            ]
            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages
            )
            solution_info = response.choices[0].message.content.strip()

            if predicted_class == "No_tomato_leaf" :
                st.write ( f"### {t['invalid_image']}" )
                st.write ( f" {t['upload_valid_image']}  " )

            elif 85 >= confidence >= 40 :
                st.write (f"### {t['invalid_image']}" )
                st.write (f"{t['upload_valid_image']}  ")

            elif 39 >= confidence >= 0 :
                st.write ( t["low_quality"] )
                st.write ( t ["upload_clear"] )

            else :
                st.write ( f"### {t['disease_detected']} {t[predicted_class]}" )
                st.write ( f"{t['confidence']} {confidence:.2f}%" )
                st.write ( t["solution_info"] )
                st.write(solution_info)

                # Generate Word Report
                if st.button ( t["generate_doc"] ) :
                    doc = Document ()
                    doc.add_heading ( t["title_doc"], level=1 )

                    # Add Disease Name
                    doc.add_paragraph ( f"**{t['disease_detected']}** {predicted_class}" )

                    # Add Image
                    if image :
                        temp_image_path = tempfile.NamedTemporaryFile ( delete=False, suffix=".png" ).name
                        image.save ( temp_image_path )
                        doc.add_picture ( temp_image_path, width=Inches ( 4.5 ) )

                    # Add Solution Info
                    doc.add_heading ( t["solution_info"], level=2 )
                    doc.add_paragraph ( solution_info )

                    # Save Word File
                    temp_doc_path = tempfile.NamedTemporaryFile ( delete=False, suffix=".docx" ).name
                    doc.save ( temp_doc_path )

                    with open ( temp_doc_path, "rb" ) as doc_file :
                        st.download_button ( t["download_report"], doc_file, "Diagnosis_Report.docx" )

        except Exception as e:
            st.error(f"Error: {e}")

    else:
        st.write(t["upload_valid_image"])

    st.markdown ( "---" )
    st.header ( t["offer"] )
    col1, col2 = st.columns ( [2, 1] )
    with col1 :
        st.subheader ( t["buy_seeds"] )
        st.write ( t["seed_benefits"] )
        st.write ( t["why_choose"] )
        st.write ( f"### {t['shop']}" )

    with col2 :
        ad_image = "./image/tomato_seeds.png"
        st.image ( ad_image, caption=t["buy_seeds"], use_container_width=True )

    # Footer
    st.markdown ( f"## {t['footer_title']}" )
    st.write ( t["footer_description"] )
    st.markdown ( f"#### {t['footer_closing']}" )
    st.markdown ( "---" )
